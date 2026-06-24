"""
RAG Service - Core logic (Phase 1, 3, 4 from the team's notebook)
This file contains all the RAG classes extracted from the notebook as-is.
"""

from unittest import result

import faiss
import numpy as np
from pathlib import Path
import os

import pdfplumber
import docx
from bs4 import BeautifulSoup

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from sentence_transformers import SentenceTransformer

import requests


# ─────────────────────────────────────────────
# Phase 1 – Data Processing
# ─────────────────────────────────────────────

def extract_file(file_path: str):
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        pages_text = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                tables = page.extract_tables() or []
                page_content = f"\n--- PAGE {i+1} ---\n"
                if text:
                    page_content += text + "\n"
                if tables:
                    page_content += "\nTABLES:\n"
                    for table in tables:
                        for row in table:
                            if row:
                                page_content += " | ".join(str(cell or "") for cell in row) + "\n"
                pages_text.append(page_content)
        full_text = "\n".join(pages_text)

    elif ext == ".docx":
        doc = docx.Document(file_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)

    elif ext == ".html":
        with open(file_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f, "lxml")
        for tag in soup(["script", "style", "nav", "header", "footer"]):
            tag.decompose()
        body = soup.find("body")
        full_text = body.get_text(separator=" ") if body else soup.get_text(separator=" ")
        full_text = " ".join(full_text.split())

    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            full_text = f.read()

    else:
        return None

    if not full_text.strip():
        print(f"Empty file: {file_path}")
        return None

    return {
        "text": full_text,
        "metadata": {"file": file_path, "type": ext}
    }


def load_documents(folder_path: str):
    documents = []
    for file in os.listdir(folder_path):
        path = os.path.join(folder_path, file)
        if not os.path.isfile(path):
            continue
        print("Processing:", file)
        result = extract_file(path)
        if result is None:
            continue
        documents.append(
            Document(page_content=result["text"], metadata=result["metadata"])
        )
    return documents


def clean_text(text: str) -> str:
    text = text.replace("\n", " ").replace("\t", " ")
    return " ".join(text.split())


def process_files(file_content, chunk_size=500, overlap=50):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap
    )
    return splitter.split_documents(file_content)


# ─────────────────────────────────────────────
# Embedding Client
# ─────────────────────────────────────────────

class EmbeddingClient:
    def __init__(self):
        self.model = SentenceTransformer("all-MiniLM-L6-v2")
        self.embedding_size = self.model.get_sentence_embedding_dimension()

    def embed(self, text, doc_type="passage"):
        return self.model.encode(text).astype("float32")


# ─────────────────────────────────────────────
# FAISS Vector Database
# ─────────────────────────────────────────────

class FAISSVectorDB:
    def __init__(self, dim):
        self.index = faiss.IndexFlatL2(dim)
        self.texts = []
        self.metadata = []

    def is_collection_exists(self, name):
        return True

    def create_collection(self, name, dim):
        self.index = faiss.IndexFlatL2(dim)
        self.texts = []
        self.metadata = []

    def add_documents(self, collection_name, texts, vectors, metadata):
        vectors = np.array(vectors).astype("float32")
        self.index.add(vectors)
        self.texts.extend(texts)
        self.metadata.extend(metadata)

    def search_by_vector(self, collection_name, query_vector, top_k=5):
        D, I = self.index.search(
            np.array([query_vector]).astype("float32"),
            top_k
        )
        results = []
        for i in I[0]:
            if i < len(self.texts):
                results.append({
                    "text": self.texts[i],
                    "metadata": self.metadata[i]
                })
        return results


# ─────────────────────────────────────────────
# Vector Pipeline
# ─────────────────────────────────────────────

class VectorPipeline:
    def __init__(self, embedding_client, vectordb_client):
        self.embedding_client = embedding_client
        self.vectordb_client = vectordb_client

    def push_data_to_index(self, project_id, chunks):
        texts = [c.page_content for c in chunks]
        metadata = [c.metadata for c in chunks]
        vectors = [self.embedding_client.embed(t) for t in texts]

        if not self.vectordb_client.is_collection_exists(project_id):
            self.vectordb_client.create_collection(
                project_id,
                self.embedding_client.embedding_size
            )

        self.vectordb_client.add_documents(project_id, texts, vectors, metadata)


def run_pipeline(folder_path, project_id, embedding_client, vectordb_client):
    print("Starting pipeline...")
    docs = load_documents(folder_path)
    if not docs:
        print("No valid files found")
        return False
    chunks = process_files(docs)
    print(f"Total chunks: {len(chunks)}")
    pipeline = VectorPipeline(embedding_client, vectordb_client)
    pipeline.push_data_to_index(project_id, chunks)
    print("Offline pipeline done")
    return True


# ─────────────────────────────────────────────
# Phase 3 – LLM & RAG System
# ─────────────────────────────────────────────

class prompt_template:
    def __init__(self):
        self.system_prompt = """
    You are a Retrieval-Augmented Generation (RAG) assistant.
    Your task is to answer ONLY using the provided context.
    """
        self.prompt_footer = """
    ### INSTRUCTIONS ###
    - Use ONLY the information present in the CONTEXT.
    - Do NOT use any external knowledge.
    - Quote exact phrases from the context when possible.
    - Do not paraphrase critical facts like job titles or locations.
    - Do NOT mention any candidate, company, or role that is NOT explicitly found in the context.
    - Do NOT invent or assume missing details.
    - If the answer is not found in the context, respond EXACTLY with:
    "I am sorry, but I do not have this information in the provided files."
    - If multiple candidates match, list ONLY those found in the context.
    - Be concise and accurate.
    """

    def rag_prompt(self, context_text, query):
        return f"""
        ### SYSTEM ###
        {self.system_prompt}

        ### CONTEXT ###
        {context_text}

        ### USER QUESTION ###
        {query}

        {self.prompt_footer}

        ### FINAL ANSWER ###
        """

class LLM_generation:
    def __init__(self):
        self.api_key = os.getenv("OPENROUTER_API_KEY", "")

    def construct_prompt(self, query, role):
        mapped_role = "assistant" if role == "model" else role
        return {"role": mapped_role, "content": query}

    def generate_response(self, full_prompt, chat_history):
        if chat_history is None:
            chat_history = []

        clean_history = [
            msg for msg in chat_history
            if isinstance(msg, dict) and "role" in msg and "content" in msg
        ]

        messages = clean_history + [{"role": "user", "content": full_prompt}]

        result = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers={"Authorization": f"Bearer {self.api_key}"},
            json={
                "model": "openrouter/free",
                "messages": messages,
                "temperature": 0.1,
                "top_p": 0.95
            }
        ).json()

        print("OpenRouter response:", result)

        if "choices" not in result:
            raise Exception(f"OpenRouter error: {result}")

        return result["choices"][0]["message"]["content"]


class Rag_system:
    def __init__(self, embedding_client_instance, faiss_db_instance, llm_generation_instance):
        self.embedding_client = embedding_client_instance
        self.faiss_db_client = faiss_db_instance
        self.generation_client = llm_generation_instance

    def search_by_vector(self, project, vector, top_k):
        return self.faiss_db_client.search_by_vector(project, vector, top_k)

    def answer_rag_question(self, project, user_query, top_k=5, chat_history=None):
        user_query_vector = self.embedding_client.embed(user_query)
        relevant_docs = self.search_by_vector(project, user_query_vector, top_k=top_k)
        context = "\n".join([doc["text"] for doc in relevant_docs])

        promptGenerator = prompt_template()
        full_prompt = promptGenerator.rag_prompt(context, user_query)

        if chat_history is None:
            chat_history = [{"role": "system", "content": promptGenerator.system_prompt}]

        answer = self.generation_client.generate_response(full_prompt, chat_history)

        chat_history.append(self.generation_client.construct_prompt(user_query, role="user"))
        chat_history.append(self.generation_client.construct_prompt(answer, role="model"))

        return answer, full_prompt, chat_history
